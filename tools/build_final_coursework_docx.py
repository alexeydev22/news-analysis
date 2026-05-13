from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = ROOT / "docs" / "final"
ASSETS_DIR = FINAL_DIR / "assets"
SOURCE_DIR = FINAL_DIR / "source"
DOCX_PATH = FINAL_DIR / "coursework-explanatory-note.docx"
TMP_ASSETS_DIR = Path("/private/tmp/coursework-docx-assets")

BLUE = "1F3A5F"
ACCENT = "0F5132"
INK = "111111"
MUTED = "4F5B57"
GRID = "B7C6C0"
TABLE_HEADER = "E8EEF5"
LIGHT_GREEN = "EAF4EF"


def load_materials_facts() -> dict[str, Any]:
    facts_path = SOURCE_DIR / "materials-facts.json"
    return json.loads(facts_path.read_text(encoding="utf-8"))


def format_number(value: int) -> str:
    return f"{value:,}".replace(",", " ")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "Arial Bold.ttf" if bold else "Arial.ttf"
    return ImageFont.truetype(f"/System/Library/Fonts/Supplemental/{name}", size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= width:
            current = candidate
            continue
        if current:
            lines.append(current)
        current = word
    if current:
        lines.append(current)
    return lines


def rounded_rect(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: str,
    outline: str | None = None,
    radius: int = 16,
    width: int = 2,
) -> None:
    fill = fill if fill.startswith("#") else f"#{fill}"
    outline = None if outline is None else outline if outline.startswith("#") else f"#{outline}"
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def build_architecture_diagram() -> Path:
    TMP_ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    path = TMP_ASSETS_DIR / "architecture-diagram.png"
    img = Image.new("RGB", (1800, 1050), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = font(46, True)
    label_font = font(28, True)
    body_font = font(22)
    small_font = font(18)

    draw.text((70, 50), "Компактная микросервисная архитектура системы", fill=f"#{INK}", font=title_font)
    draw.text(
        (72, 112),
        "Пользовательский интерфейс, API-слой, сервисы анализа, поиска, диалога и инфраструктурные хранилища",
        fill=f"#{MUTED}",
        font=body_font,
    )

    boxes = {
        "ui": (70, 220, 360, 360, "React UI", "чат, отчет, прогноз"),
        "news": (70, 520, 360, 660, "news-service", "новости и датасет"),
        "gateway": (500, 260, 820, 400, "api-gateway", "единый API и SSE"),
        "worker": (500, 560, 820, 700, "news-worker", "фоновая индексация"),
        "retrieval": (960, 170, 1270, 310, "retrieval-service", "поиск источников"),
        "analysis": (960, 390, 1270, 530, "analysis-service", "ML-анализ влияния"),
        "dialog": (960, 610, 1270, 750, "dialog-service", "RAG-ответ"),
        "qdrant": (1430, 170, 1710, 310, "Qdrant", "векторный индекс"),
        "redis": (1430, 450, 1710, 590, "Redis", "очередь и события"),
        "mlflow": (1430, 730, 1710, 870, "MLflow", "метрики моделей"),
    }

    for key, (x1, y1, x2, y2, title, subtitle) in boxes.items():
        fill = "EAF4EF" if key in {"ui", "gateway"} else "F8FAF9"
        outline = ACCENT if key in {"ui", "gateway"} else GRID
        rounded_rect(draw, (x1, y1, x2, y2), fill, outline, radius=18, width=3)
        draw.text((x1 + 24, y1 + 28), title, fill=f"#{INK}", font=label_font)
        draw.text((x1 + 24, y1 + 78), subtitle, fill=f"#{MUTED}", font=body_font)

    def arrow(points: list[tuple[int, int]], label: str) -> None:
        draw.line(points, fill=f"#{ACCENT}", width=4, joint="curve")
        sx, sy = points[-2]
        ex, ey = points[-1]
        if abs(ex - sx) >= abs(ey - sy) and ex >= sx:
            pts = [(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)]
        elif abs(ex - sx) >= abs(ey - sy):
            pts = [(ex, ey), (ex + 16, ey - 9), (ex + 16, ey + 9)]
        elif ey >= sy:
            pts = [(ex, ey), (ex - 9, ey - 16), (ex + 9, ey - 16)]
        else:
            pts = [(ex, ey), (ex - 9, ey + 16), (ex + 9, ey + 16)]
        draw.polygon(pts, fill=f"#{ACCENT}")
        mx, my = points[len(points) // 2]
        draw.text((mx - 45, my - 28), label, fill=f"#{MUTED}", font=small_font)

    arrow([(360, 290), (500, 330)], "запрос")
    arrow([(215, 360), (215, 520)], "данные")
    arrow([(360, 590), (500, 630)], "задачи")
    arrow([(820, 310), (960, 240)], "поиск")
    arrow([(820, 330), (960, 460)], "анализ")
    arrow([(820, 350), (900, 350), (900, 680), (960, 680)], "ответ")
    arrow([(1270, 240), (1430, 240)], "vectors")
    arrow([(820, 630), (900, 630), (900, 300), (960, 240)], "index")
    arrow([(820, 660), (1430, 520)], "queue")
    arrow([(1270, 460), (1350, 460), (1350, 800), (1430, 800)], "metrics")
    draw.text(
        (70, 940),
        "Архитектура оставляет backend компактным: каждый сервис закрывает одну прикладную роль ML/RAG-сценария.",
        fill=f"#{MUTED}",
        font=body_font,
    )
    img.save(path)
    return path


def build_rag_pipeline_diagram() -> Path:
    TMP_ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    path = TMP_ASSETS_DIR / "rag-pipeline-diagram.png"
    img = Image.new("RGB", (1800, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = font(44, True)
    label_font = font(27, True)
    body_font = font(21)
    small_font = font(18)

    draw.text((70, 48), "ML/RAG-пайплайн анализа экономических новостей", fill=f"#{INK}", font=title_font)
    steps = [
        ("Вопрос", "пользователь формулирует экономический запрос"),
        ("Retrieval", "поиск релевантных новостей в Qdrant"),
        ("ML-классификация", "оценка влияния: positive, neutral, negative"),
        ("Контекст", "отбор источников, score и краткие признаки"),
        ("Внешний LLM API", "аналитическое обобщение по найденному контексту"),
        ("Ответ", "вывод ответа, источников и прогноза в UI"),
    ]

    x = 70
    y = 230
    w = 235
    h = 150
    gap = 58
    for idx, (title, text) in enumerate(steps):
        fill = LIGHT_GREEN if idx in {0, 5} else "F8FAF9"
        rounded_rect(draw, (x, y, x + w, y + h), fill, ACCENT if idx in {0, 5} else GRID, radius=18, width=3)
        draw.text((x + 22, y + 24), title, fill=f"#{INK}", font=label_font)
        yy = y + 70
        for line in wrap_text(draw, text, body_font, w - 44)[:3]:
            draw.text((x + 22, yy), line, fill=f"#{MUTED}", font=body_font)
            yy += 28
        if idx < len(steps) - 1:
            ax = x + w + 8
            draw.line((ax, y + h // 2, ax + gap - 16, y + h // 2), fill=f"#{ACCENT}", width=4)
            draw.polygon(
                [(ax + gap - 16, y + h // 2), (ax + gap - 32, y + h // 2 - 10), (ax + gap - 32, y + h // 2 + 10)],
                fill=f"#{ACCENT}",
            )
        x += w + gap

    lanes = [
        ("Данные", "FNSPID, 50 000 новостей, подготовка текста, train/validation/test split"),
        ("Модели", "TF-IDF + Logistic Regression; sentence embeddings + Logistic Regression; lightweight transformer"),
        ("Метрики", "accuracy показывает долю верных ответов, macro F1 устойчивее отражает качество по всем классам"),
    ]
    y = 530
    for title, text in lanes:
        rounded_rect(draw, (90, y, 1710, y + 82), "F8FAF9", GRID, radius=14, width=2)
        draw.text((120, y + 22), title, fill=f"#{INK}", font=label_font)
        draw.text((330, y + 26), text, fill=f"#{MUTED}", font=small_font)
        y += 105

    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 120, start: int = 120, bottom: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size: int = 14, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11) -> None:
    cell.text = ""
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Cm(width)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True, size=10)
        set_cell_shading(table.rows[0].cells[index], TABLE_HEADER)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=10)
    doc.add_paragraph()


def add_numbers(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(item)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(text)
    style_run(run, size=12)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_section_title(doc: Document, title: str, level: int = 1) -> None:
    doc.add_heading(title, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def configure_doc(doc: Document, facts: dict[str, Any]) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after, align in [
        ("Title", 16, INK, 0, 12, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 1", 14, INK, 18, 8, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, INK, 12, 6, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 14, INK, 10, 4, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = align

    for list_style_name in ["List Bullet", "List Number"]:
        list_style = styles[list_style_name]
        list_style.font.name = "Times New Roman"
        list_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        list_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        list_style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        list_style.font.size = Pt(14)

    student = facts["student"]
    header = section.header.paragraphs[0]
    header.text = "Пояснительная записка к курсовой работе"
    header.runs[0].font.name = "Times New Roman"
    header.runs[0].font.size = Pt(10)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = f"{student['name']}, {student['group']}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.name = "Times New Roman"
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_title_page(doc: Document, facts: dict[str, Any]) -> None:
    student = facts["student"]
    coursework = facts["coursework"]
    rows = [
        student["university"],
        student["department"],
    ]
    for row in rows:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(row)
        style_run(run, size=14, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    for text, size in [
        ("КУРСОВАЯ РАБОТА", 16),
        (f"по дисциплине: «{coursework['discipline']}»", 14),
        (f"на тему: «{coursework['topic']}»", 14),
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(text)
        style_run(run, size=size, bold=True)

    for _ in range(5):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.first_line_indent = Cm(0)
    run = meta.add_run(f"Выполнил: {student['name']}\nГруппа: {student['group']}\nРуководитель: ____________________")
    style_run(run, size=14)

    for _ in range(2):
        doc.add_paragraph()

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year.paragraph_format.first_line_indent = Cm(0)
    run = year.add_run(f"{student['city']}, {student['year']}")
    style_run(run, size=14)
    add_page_break(doc)


def add_contents(doc: Document) -> None:
    add_section_title(doc, "СОДЕРЖАНИЕ")
    entries = [
        "Введение",
        "1 Теоретические основы анализа экономических новостей",
        "2 Данные и постановка ML-задачи",
        "3 Методы и модели",
        "4 Архитектура программной системы",
        "5 Реализация и пользовательский сценарий",
        "6 Эксперимент и результаты",
        "7 Направления улучшения",
        "Заключение",
        "Список использованных источников",
        "Приложения",
    ]
    for entry in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.add_run(entry)
    add_page_break(doc)


def intro_paragraphs(facts: dict[str, Any]) -> list[str]:
    rows = facts["dataset"]["rows"]
    return [
        "Экономические новости являются одним из основных источников оперативной информации о состоянии рынков, деловой активности, инфляционных ожиданий, занятости, торговых потоков и решений регуляторов. Для отдельного пользователя, аналитика или учебного исследователя поток таких сообщений становится слишком большим для ручного просмотра. При этом ценность имеет не только факт обнаружения новости, но и ее интерпретация: относится ли сообщение к положительному, отрицательному или нейтральному влиянию, какие источники подтверждают вывод и как кратко объяснить возможное значение события.",
        "Актуальность работы определяется ростом объема текстовых экономических данных и необходимостью объединять поиск, машинное обучение и языковую модель в одном прикладном сценарии. Простая выдача списка документов не решает задачу анализа, поскольку пользователь ожидает связный ответ на естественном языке. Обычная генерация без источников также недостаточна, так как результат должен опираться на найденные материалы. Поэтому в работе рассматривается автоматическая диалоговая аналитическая система, которая сначала извлекает релевантный контекст из корпуса новостей, затем применяет классификаторы влияния и только после этого формирует ответ.",
        f"В качестве данных используется набор FNSPID объемом {format_number(rows)} экономических новостей. В рамках проекта решается задача классификации влияния новости по трем классам: positive, neutral и negative. Такая постановка связана с прикладным вопросом пользователя: не просто найти текст, а понять направленность события и использовать этот результат при формировании аналитического ответа. Дополнительно система поддерживает тематический прогноз по группам новостей и формирует ML-отчет с метриками моделей.",
        "Объектом исследования являются экономические новости как источник текстовых данных для семантического анализа. Предметом исследования являются методы автоматической обработки этих новостей: текстовое представление, классификация влияния, поиск релевантных документов, retrieval-augmented generation и интеграция внешнего LLM API в пользовательский сценарий.",
        "Цель курсовой работы состоит в разработке автоматической диалоговой системы для анализа экономических новостей, которая принимает вопрос пользователя, находит релевантные документы, оценивает влияние новостей, формирует ответ на основе источников и отображает результаты в веб-интерфейсе.",
        "Для достижения цели поставлены задачи: изучить теоретические основы NLP и анализа экономических новостей; подготовить датасет и формализовать ML-задачу; реализовать и сравнить несколько моделей классификации; спроектировать компактную backend-архитектуру; реализовать сценарий поиска, анализа и ответа; провести эксперимент, интерпретировать метрики и определить направления улучшения качества.",
        "Практическая значимость работы состоит в том, что полученная система показывает полный путь от набора данных до пользовательского аналитического продукта. Она демонстрирует, как учебная ML-задача классификации может быть встроена в реальный диалоговый интерфейс, где важны не только метрики модели, но и воспроизводимость обработки, объяснимость результата, наличие источников и удобство просмотра отчета.",
    ]


def theory_paragraphs() -> list[str]:
    return [
        "Анализ экономических новостей относится к задачам обработки естественного языка, в которых текстовая информация преобразуется в признаки и используется для принятия аналитического решения. В отличие от числовых временных рядов, новостные данные содержат неструктурированные формулировки, контекст, неоднозначность и различный уровень детализации. Поэтому первым этапом становится приведение текста к форме, пригодной для поиска и классификации.",
        "Классический подход к представлению текста строится на статистических признаках, например TF-IDF. Метод отражает важность слова или n-граммы в конкретном документе с учетом частоты в корпусе. Для экономических новостей такой подход полезен тем, что отдельные термины часто напрямую связаны с характером события: инфляция, рост, сокращение, ставка, безработица, экспорт, прибыль, дефицит. Модель на TF-IDF-признаках остается объяснимой и хорошо подходит как baseline.",
        "Семантические embedding-представления решают другую задачу: они переводят текст в плотный вектор, где близость отражает не только буквальное совпадение слов, но и смысловую близость сообщений. Это особенно важно для retrieval-сценария, где пользователь может задать вопрос словами, отличающимися от текста новости. Например, запрос о снижении ценового давления должен находить новости об инфляции даже без полного совпадения формулировок.",
        "Классификация влияния новостей представляет собой задачу supervised learning. Каждому тексту сопоставляется целевой класс: positive, neutral или negative. Positive означает потенциально благоприятное влияние на экономическую ситуацию или ожидания; negative указывает на риск или ухудшение; neutral фиксирует отсутствие явной направленности либо сбалансированное сообщение. В практическом применении такая классификация помогает не только упорядочить источники, но и объяснить итоговую позицию системы.",
        "Retrieval-augmented generation объединяет поиск и генерацию ответа. Сначала система извлекает релевантные документы, затем передает их как контекст генератору ответа. В результате языковая модель не отвечает в отрыве от данных, а опирается на найденные новости. В проекте этот принцип применяется к экономическому домену: пользователь получает краткий аналитический ответ, а рядом видит источники, их релевантность и оценку влияния.",
        "Языковые модели в аналитических системах выполняют роль компонента обобщения, а не единственного источника истины. Внешний LLM API используется после поиска и ML-анализа, поэтому его ответ ограничен переданным контекстом. Такой дизайн снижает риск нерелевантной генерации и соответствует требованию объяснимости: пользователь может проверить, какие новости были использованы при формировании вывода.",
        "Для оценки качества классификаторов применяются accuracy и macro F1. Accuracy показывает долю верно классифицированных примеров, но при дисбалансе классов может переоценивать модель, хорошо распознающую только самый частый класс. Macro F1 усредняет качество по классам без учета их частоты, поэтому лучше показывает, насколько модель справляется с minor-классами и сохраняет полезность для анализа неоднородного новостного потока.",
    ]


def data_paragraphs(facts: dict[str, Any]) -> list[str]:
    rows = facts["dataset"]["rows"]
    classes = ", ".join(facts["dataset"]["classes"])
    return [
        f"Исходной базой исследования является датасет {facts['dataset']['name']}, содержащий {format_number(rows)} экономических новостей. Для курсовой работы этот объем достаточен, чтобы обучить несколько моделей, получить устойчивые метрики и проверить работу поиска на реальном корпусе, а не на искусственном наборе демонстрационных фраз.",
        "Каждая новость рассматривается как текстовый документ, в котором присутствуют заголовок, содержание или объединенное текстовое поле, а также метка влияния. В прикладной системе эти документы используются дважды: как обучающие примеры для классификаторов и как корпус для retrieval-поиска. Поэтому качество подготовки текста влияет одновременно на ML-метрики и на пользовательское качество ответа.",
        f"Целевая переменная принимает три значения: {classes}. Наличие трех классов делает задачу более содержательной, чем бинарная оценка. Нейтральные новости важны, потому что в экономическом потоке значительная часть сообщений описывает факты без однозначной рыночной направленности. Если игнорировать этот класс, система будет искусственно усиливать позитивную или негативную интерпретацию.",
        "Перед обучением выполняется базовая подготовка текста: удаление пустых значений, приведение записей к единой схеме, проверка целевой метки, формирование обучающего и тестового представления. Для моделей на TF-IDF важна стабильность токенизации и ограничение словаря, чтобы случайные редкие слова не перегружали признаковое пространство. Для embedding-подхода важна сохранность семантического содержания фразы.",
        "Данные делятся на train, validation и test. Обучающая часть используется для подбора параметров модели, validation помогает контролировать качество в процессе настройки, test применяется только для итогового сравнения. Такое разделение необходимо, чтобы метрики отражали обобщающую способность модели, а не запоминание конкретных новостей.",
        "Особое внимание уделяется дисбалансу классов. В новостных данных одни типы сообщений встречаются чаще других: нейтральные публикации могут доминировать над явно позитивными или негативными. В такой ситуации модель может показывать приемлемую accuracy, но плохо распознавать редкие классы. Поэтому в отчете рядом с accuracy обязательно анализируется macro F1.",
        "С точки зрения пользовательского сценария датасет используется не изолированно, а как активный источник знаний. После индексации новости становятся доступными для поиска по вопросу пользователя. Это связывает исследовательскую часть с программной реализацией: корпус, на котором обучается классификатор, одновременно является базой фактов для RAG-ответа и тематического прогноза.",
        "Для экономических новостей особенно важна воспроизводимость подготовки корпуса. Если одна и та же новость в разных экспериментах получает различное представление, сравнение моделей теряет смысл. Поэтому подготовка данных должна быть детерминированной: одинаковые правила очистки, одинаковое разбиение, одинаковая интерпретация целевой колонки и одинаковые ограничения на признаки.",
        "Еще одна особенность датасета состоит в том, что экономическая направленность часто выражается не одним словом, а сочетанием показателя и контекста. Например, рост инфляции может восприниматься негативно, но рост выручки компании обычно положителен. Это означает, что модель должна учитывать локальные словосочетания, а не только отдельные токены. Поэтому для TF-IDF полезны n-граммы, а для transformer-подхода важен достаточно длинный контекст.",
        "В прикладном RAG-сценарии качество данных проявляется и через источники, которые видит пользователь. Если корпус содержит дубликаты, пустые тексты или нерелевантные записи, retrieval может вернуть слабый контекст даже при хорошей модели классификации. Поэтому работа с данными рассматривается как отдельная часть проекта, а не как вспомогательная техническая операция.",
    ]


def methods_paragraphs() -> list[str]:
    return [
        "В работе сравниваются три подхода к классификации влияния экономических новостей. Такое сравнение выбрано намеренно: оно показывает разницу между классическим статистическим baseline, семантическими embedding-признаками и легкой нейросетевой моделью. Для курсовой работы важно не только получить лучший результат, но и объяснить, почему разные методы ведут себя по-разному на одном наборе данных.",
        "Первая модель, TF-IDF + Logistic Regression, служит основным baseline. TF-IDF преобразует текст в разреженный вектор признаков, а логистическая регрессия обучает линейную границу между классами. Преимущество подхода состоит в скорости, воспроизводимости и интерпретируемости. Если модель показывает высокое качество, это означает, что в данных есть достаточно сильные лексические сигналы, связанные с классом влияния.",
        "Вторая модель, embedding-logreg, использует sentence embeddings как признаки для логистической регрессии. Она должна лучше учитывать смысловую близость текстов, но результат зависит от того, насколько embedding-модель соответствует языку, домену и стилю новостей. На текущей разметке embedding-подход оказался слабее TF-IDF, что указывает на важность доменных слов и явных лексических маркеров в задаче влияния.",
        "Третья модель, tiny-transformer-classifier, представляет легкий transformer-классификатор. Такой вариант включен для проверки нейросетевого подхода, однако маленький размер модели и ограниченный режим обучения сдерживают качество. Transformer-модели обычно требуют более аккуратной настройки, достаточного времени обучения, баланса классов и стабильной разметки, иначе они могут проиграть простому baseline.",
        "Все модели оцениваются на одном тестовом разбиении, что делает сравнение корректным. Для каждой модели фиксируются test accuracy и test macro F1. Accuracy удобна для общей оценки, а macro F1 показывает, насколько равномерно модель работает по всем классам. Это особенно важно для экономических новостей, где ошибки на редком негативном классе могут быть прикладно значимее, чем ошибки на частом нейтральном классе.",
        "Выбор нескольких моделей также полезен для архитектуры системы. Пользовательский интерфейс и API не должны зависеть от конкретного классификатора. Если в будущем будет обучена более сильная модель, ее можно подключить в analysis-service без изменения сценария чата, отчета и прогноза. Такое разделение повышает устойчивость проекта как программной системы.",
        "TF-IDF baseline также удобен для анализа ошибок. После обучения можно посмотреть, какие слова и словосочетания сильнее всего связаны с каждым классом. Это помогает понять, действительно ли модель выучила экономически осмысленные признаки или опирается на случайные артефакты корпуса. Для учебной работы такая объяснимость важна не меньше абсолютного значения метрики.",
        "Embedding-подход полезен как проверка гипотезы о смысловой близости. Если две новости описывают одно и то же экономическое явление разными словами, embedding должен расположить их рядом в векторном пространстве. Даже если итоговая классификация оказалась слабее TF-IDF, этот подход остается значимым для retrieval и дальнейшего развития поиска.",
        "Transformer-классификатор включен как исследовательский ориентир. Он показывает, что архитектура проекта допускает более сложные модели, но одновременно демонстрирует необходимость аккуратной экспериментальной дисциплины. Нейросетевая модель требует контроля переобучения, качества токенизации, длины входа, баланса классов и достаточного количества эпох.",
        "Сравнение моделей не ограничивается выбором победителя. Оно формирует практическое решение для текущей версии системы: использовать лучший baseline в пользовательском сценарии, сохранить альтернативные подходы в отчете и зафиксировать, какие условия нужны для повышения качества в следующих итерациях.",
    ]


def architecture_paragraphs() -> list[str]:
    return [
        "Программная система построена как набор связанных сервисов, каждый из которых отвечает за отдельный этап аналитического сценария. Архитектура описывается компактно, поскольку основной фокус работы находится на данных, моделях и результатах эксперимента. Тем не менее разделение backend-компонентов важно: оно показывает, как ML-логика встроена в приложение, а не остается отдельным исследовательским скриптом.",
        "Frontend на React предоставляет пользователю страницы чата, ML-отчета и тематического прогноза. API Gateway принимает запросы от интерфейса и координирует обращение к сервисам анализа, поиска и диалога. News-service отвечает за корпус новостей и подготовку данных для индексации. News-worker выполняет фоновые задачи, чтобы тяжелая обработка не блокировала пользовательский запрос.",
        "Retrieval-service отвечает за поиск релевантных документов в векторном хранилище Qdrant. Сервис возвращает новости, которые по смыслу ближе всего к вопросу пользователя. Analysis-service применяет ML-модель и определяет влияние найденных новостей. Dialog-service получает вопрос, контекст и результаты анализа, после чего формирует итоговый ответ, используя внешний LLM API или резервный шаблонный режим.",
        "Redis применяется для очередей и событий, Qdrant хранит векторный индекс новостей, MLflow фиксирует результаты экспериментов и артефакты моделей. Такой набор инфраструктурных компонентов достаточен для учебного проекта: он демонстрирует production-подход, но не перегружает курсовую работу деталями промышленной эксплуатации.",
        "Ключевой архитектурный принцип состоит в том, что внешняя языковая модель не заменяет ML-часть. Она получает уже найденный и предварительно проанализированный контекст. Поэтому система остается объяснимой: можно отдельно проверить корпус данных, retrieval-результаты, классификацию влияния и итоговую генерацию.",
        "Границы сервисов выбраны так, чтобы каждая часть могла развиваться независимо. Например, retrieval-service может менять способ построения embeddings и параметры поиска без изменения интерфейса чата. Analysis-service может подключить новый классификатор, сохранив прежний формат ответа. Dialog-service может изменить стратегию промпта или поставщика внешнего LLM API, не затрагивая подготовку датасета и ML-отчет.",
        "Компактность backend-описания не означает упрощения ML-сценария. Напротив, архитектура нужна для того, чтобы экспериментальная часть была доступна пользователю через приложение: результаты моделей отображаются в отчете, найденные источники выводятся в чате, а прогноз строится по индексированному корпусу. Поэтому программная структура поддерживает исследовательскую логику курсовой работы.",
        "Такой подход также упрощает диагностику ошибок. Если пользователь получает слабый ответ, можно последовательно проверить этапы: был ли корректно подготовлен корпус, нашел ли retrieval релевантные документы, правильно ли analysis-service оценил влияние, достаточно ли контекста получил dialog-service. Разделение компонентов делает качество системы наблюдаемым.",
    ]


def realization_paragraphs() -> list[str]:
    return [
        "Пользовательский сценарий начинается с естественного вопроса об экономической ситуации. Пользователь не обязан знать структуру датасета или формулировать запрос в виде ключевых слов. Система принимает текст вопроса, передает его в backend и запускает цепочку retrieval, ML-анализа и генерации ответа.",
        "На этапе поиска retrieval-service извлекает несколько наиболее релевантных новостей. Для каждой найденной записи сохраняются заголовок или краткое описание, источник, score релевантности и текстовый фрагмент. Эти данные нужны не только для ответа, но и для интерфейса: пользователь должен видеть, на каких материалах основан вывод.",
        "После поиска analysis-service применяет классификатор влияния. Результат включает класс positive, neutral или negative, численную уверенность и краткое объяснение. Эти признаки помогают dialog-service сформировать не абстрактный пересказ, а аналитический ответ с учетом направленности найденных новостей.",
        "Dialog-service собирает промпт для внешнего LLM API на основе вопроса, найденных документов и оценок влияния. Важным ограничением является использование только переданного контекста: ответ должен объяснять найденные новости, а не подменять их внешними рассуждениями. Если внешний API недоступен, система может использовать резервный шаблонный ответ, сохраняя работоспособность базового сценария.",
        "Страница чата показывает центральный результат проекта: вопрос, ответ, список источников и статусы обработки. Это делает ML/RAG-процесс видимым для пользователя. Вместо обычного чат-бота реализована аналитическая система, где ответ связан с корпусом новостей и сопровождается проверяемыми источниками.",
        "Страница ML-отчета предназначена для контроля исследовательской части. В ней отображаются размер датасета, список обученных моделей, метрики качества и интерпретация результатов. Наличие такой страницы важно для защиты, потому что показывает, что система основана на сравнении моделей, а не только на интерфейсной демонстрации.",
        "Страница прогноза показывает агрегированный анализ тем. В текущих фактах проекта отражено, что прогноз строится по 10 000 документам и 2 716 темам. Этот сценарий расширяет обычный поиск: пользователь видит не только ответ на отдельный вопрос, но и тематическое обобщение корпуса новостей.",
        "В реализации важно, что пользовательский интерфейс не скрывает исследовательскую природу проекта. Чат демонстрирует конечный сценарий, ML-отчет показывает качество моделей, а прогноз показывает агрегированную работу с корпусом. Эти три страницы вместе образуют полный контур: данные, модель, поиск, ответ и интерпретация результата.",
        "Сценарий работы с источниками также важен для доверия к системе. Пользователь может сопоставить итоговый ответ с найденными карточками новостей и увидеть, какие документы повлияли на вывод. Для экономической аналитики это принципиально: ответ без проверяемых источников сложнее использовать и сложнее защищать как результат машинного обучения.",
        "ML-отчет выполняет роль внутреннего контроля качества. Если в дальнейшем изменится датасет, признаки или модель, отчет позволит быстро увидеть, улучшилось ли качество на test split или изменения только усложнили систему без прироста метрик. Поэтому отчет встроен в приложение как часть жизненного цикла модели.",
    ]


def experiment_paragraphs(facts: dict[str, Any]) -> list[str]:
    best = facts["models"][0]
    return [
        "Экспериментальная часть направлена на сравнение трех моделей классификации влияния новостей. Все модели решают одну и ту же задачу, используют общий набор классов и оцениваются по test accuracy и test macro F1. Такой дизайн эксперимента позволяет сопоставить методы без смешения факторов, связанных с разными данными или разными метриками.",
        f"Лучший результат показала модель {best['name']} с test accuracy {best['test_accuracy']:.3f} и test macro F1 {best['test_macro_f1']:.3f}. Это означает, что классический TF-IDF baseline оказался наиболее устойчивым на текущем корпусе и разметке. Для экономических новостей такой результат объясним: многие признаки влияния выражены лексически и хорошо фиксируются частотными признаками.",
        "Embedding-logreg уступила TF-IDF. Возможная причина состоит в том, что универсальные embedding-признаки сглаживают отдельные доменные маркеры, важные для тонкой классификации влияния. В задаче, где слова о росте, падении, ставке, инфляции или риске напрямую связаны с меткой, простая разреженная модель может сохранять больше полезной информации.",
        "Tiny-transformer-classifier показала самый слабый результат среди трех моделей. Это не означает, что transformer-подход неприменим к задаче. Скорее, результат отражает ограничения маленькой модели, режима обучения и качества разметки. Для уверенного преимущества transformer-классификатора требуются более сильная модель, больше вычислительных ресурсов, настройка гиперпараметров и работа с дисбалансом классов.",
        "Macro F1 рассматривается как ключевая дополнительная метрика, потому что новостные классы могут быть распределены неравномерно. Если модель правильно классифицирует большинство нейтральных сообщений, но пропускает негативные или позитивные новости, accuracy будет выглядеть приемлемо, однако прикладная ценность снизится. Macro F1 штрафует такую ситуацию и поэтому лучше отражает качество по всем типам влияния.",
        "Полученные метрики являются умеренными, но содержательно полезными. Они показывают, что задача не сводится к простому словарному правилу и требует дальнейшей работы с данными. При этом лучший baseline уже может использоваться как компонент RAG-системы: он дает первичную оценку влияния источников, которую затем учитывает внешний LLM API при формировании ответа.",
        "Эксперимент также показал инженерную ценность ML-отчета. Пользователь или проверяющий может видеть не только итоговый ответ, но и сведения о том, какие модели были обучены и с каким качеством. Это повышает доверие к системе и делает проект ближе к полноценному аналитическому приложению.",
        "Интерпретация результатов должна учитывать учебный характер набора моделей. Цель эксперимента не состояла в том, чтобы любой ценой получить максимальную метрику. Важнее было построить воспроизводимое сравнение разных классов подходов и показать, какой вариант рационально использовать в текущей версии приложения. По этому критерию TF-IDF + Logistic Regression является обоснованным выбором.",
        "Для анализа качества важно смотреть не только на средний результат, но и на типичные ошибки. Нейтральные новости могут ошибочно относиться к позитивным или негативным, если содержат слова с яркой эмоциональной окраской, но описывают факт без явного рыночного эффекта. Обратная ситуация возникает, когда негативное событие описано формально и без сильных маркеров риска.",
        "Результаты эксперимента прямо влияют на пользовательский сценарий. Если модель ошибается, внешний LLM API может получить искаженную оценку влияния источника. Поэтому классификатор должен рассматриваться как часть общей цепочки качества: слабая метка ухудшает контекст, а слабый контекст снижает ценность итогового ответа.",
        "В текущей версии системы метрики достаточны для демонстрации полного ML/RAG-процесса, но не должны трактоваться как финальный промышленный уровень качества. Для перехода к более строгому применению необходимо расширять оценку: смотреть confusion matrix, качество по каждому классу, устойчивость на новых источниках и поведение модели при изменении новостной тематики.",
    ]


def improvements_paragraphs() -> list[str]:
    return [
        "Первое направление улучшения связано с балансировкой классов. Если один класс доминирует, модель может склоняться к нему даже при наличии признаков другого влияния. Возможные решения включают взвешивание классов, стратифицированную выборку, oversampling редких классов и более аккуратный подбор порога принятия решения.",
        "Второе направление связано с очисткой и уточнением разметки. Экономические новости часто неоднозначны: одна и та же публикация может быть позитивной для одного сектора и негативной для другого. Если метка задается слишком грубо, модель обучается на шумном сигнале. Улучшение аннотаций и введение правил для спорных случаев может заметно повысить macro F1.",
        "Третье направление состоит в подборе гиперпараметров. Для TF-IDF полезно исследовать диапазон n-грамм, размер словаря, минимальную частоту токенов и регуляризацию логистической регрессии. Для embedding-logreg можно сравнить разные embedding-модели и способы нормализации векторов. Для transformer-классификатора важны learning rate, число эпох, batch size и стратегия early stopping.",
        "Четвертое направление связано с переходом к более сильной transformer-модели. Небольшой классификатор удобен для локального стенда, но ограничен по способности извлекать сложные контекстные зависимости. Более сильная модель, дообученная на экономических новостях, может лучше распознавать тональность влияния, особенно если в тексте нет явных слов-маркеров.",
        "Пятое направление касается расширения признаков. Помимо текста новости можно использовать дату, источник, тему, макроэкономический показатель, сектор рынка и признаки похожих документов. Комбинация текстовых и структурных признаков может улучшить качество, потому что экономическое влияние часто зависит от контекста события.",
        "Шестое направление относится к тематическому прогнозированию. Текущий прогноз показывает агрегированную картину по корпусу, но его можно развивать через динамику тем во времени, сравнение периодов, оценку устойчивости к шуму и визуализацию трендов. Это сделает систему более полезной для анализа новостного потока, а не только для ответа на отдельный вопрос.",
        "Седьмое направление связано с оценкой retrieval-качества. Для RAG-системы недостаточно иметь хороший классификатор: если поиск возвращает нерелевантные новости, итоговый ответ будет слабым. Поэтому нужно отдельно измерять качество поиска, например через ручную разметку релевантности, precision@k и анализ запросов, на которых источники подбираются хуже всего.",
        "Восьмое направление относится к контролю генерации ответа. Внешний LLM API должен работать в заданных границах: использовать найденный контекст, не добавлять неподтвержденные факты и отделять аналитический вывод от источников. Для этого можно развивать промпт, добавлять автоматические проверки ответа и сохранять трассировку использованных документов.",
        "Девятое направление связано с пользовательской интерпретацией результата. Система может показывать не только итоговый класс влияния, но и распределение вероятностей, наиболее значимые признаки для baseline-модели и объяснение причин выбора источников. Это повысит прозрачность и сделает систему полезнее для учебного и аналитического применения.",
    ]


def conclusion_paragraphs() -> list[str]:
    return [
        "В результате курсовой работы разработана автоматическая диалоговая система для анализа экономических новостей. Система объединяет подготовку данных, классификацию влияния, retrieval-поиск, RAG-ответ, внешний LLM API, тематический прогноз и пользовательский интерфейс. Тем самым проект соответствует заявленной теме и демонстрирует применение машинного обучения в семантическом анализе.",
        "В работе были изучены теоретические основы NLP, классификации текстов, retrieval-augmented generation и использования языковых моделей в аналитических системах. На основе датасета FNSPID объемом 50 000 новостей сформулирована ML-задача классификации влияния по трем классам: positive, neutral и negative.",
        "Были реализованы и сравнены три модели: TF-IDF + Logistic Regression, embedding-logreg и tiny-transformer-classifier. Лучший результат показала TF-IDF + Logistic Regression, достигнув test accuracy 0.785 и test macro F1 0.717. Интерпретация результатов показала, что классический baseline хорошо отражает доменные лексические признаки текущей разметки.",
        "Архитектура программной системы построена компактно: React UI взаимодействует с API Gateway, а отдельные сервисы отвечают за новости, поиск, анализ и диалог. Qdrant используется для retrieval, Redis для очередей и событий, MLflow для фиксации экспериментов. Backend-часть поддерживает ML-сценарий, но не становится главным предметом записки.",
        "Практическим результатом является система, в которой пользователь задает вопрос, получает ответ по найденным источникам, видит страницы ML-отчета и тематического прогноза. Это отличает проект от простого чат-бота: ответ формируется на основе корпуса новостей и сопровождается проверяемым контекстом.",
        "Поставленная цель достигнута. Основные направления дальнейшего развития включают балансировку классов, уточнение разметки, подбор гиперпараметров, обучение более сильной transformer-модели, расширение признаков и развитие тематического прогнозирования.",
    ]


def add_model_results_table(doc: Document, facts: dict[str, Any]) -> None:
    interpretations = {
        "tfidf-logreg": "лучшая baseline-модель",
        "embedding-logreg": "уступает TF-IDF на текущей разметке",
        "tiny-transformer-classifier": "ограничен размером и режимом обучения",
    }
    rows = [
        [
            model["name"],
            f"{model['test_accuracy']:.3f}",
            f"{model['test_macro_f1']:.3f}",
            interpretations[model["name"]],
        ]
        for model in facts["models"]
    ]
    add_table(
        doc,
        ["Модель", "Test accuracy", "Test macro F1", "Интерпретация"],
        rows,
        widths=[4.0, 3.0, 3.0, 6.0],
    )


def add_sources(doc: Document, facts: dict[str, Any]) -> None:
    sources = facts["sources"][:9] + [
        "Lewis P. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.",
    ]
    add_section_title(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    add_numbers(doc, sources)


def add_appendices(doc: Document, facts: dict[str, Any]) -> None:
    add_page_break(doc)
    add_section_title(doc, "ПРИЛОЖЕНИЯ")
    add_section_title(doc, "Приложение А. Ссылка на репозиторий", level=2)
    doc.add_paragraph("GitHub-репозиторий проекта: https://github.com/alexeydev22/news-analysis")
    add_section_title(doc, "Приложение Б. Сводные факты проекта", level=2)
    add_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Тема", facts["coursework"]["topic"]],
            ["Датасет", facts["dataset"]["name"]],
            ["Объем корпуса", f"{format_number(facts['dataset']['rows'])} новостей"],
            ["Классы", ", ".join(facts["dataset"]["classes"])],
            ["Документы в прогнозе", format_number(facts["forecast"]["documents"])],
            ["Количество тем", format_number(facts["forecast"]["topics"])],
            ["Количество отчетов моделей", str(facts["forecast"]["model_reports"])],
        ],
        widths=[5.0, 11.0],
    )
    add_section_title(doc, "Приложение В. Роли компонентов ML/RAG-сценария", level=2)
    add_table(
        doc,
        ["Компонент", "Роль"],
        [
            ["FNSPID", "Корпус экономических новостей для обучения моделей и поиска источников."],
            ["TF-IDF + Logistic Regression", "Лучший baseline для классификации влияния на текущей разметке."],
            ["Qdrant", "Векторное хранилище для поиска релевантных новостей по вопросу."],
            ["Внешний LLM API", "Компонент обобщения найденного контекста в аналитический ответ."],
            ["MLflow", "Хранение сведений об экспериментах и метриках моделей."],
        ],
        widths=[5.0, 11.0],
    )


def build_docx() -> Path:
    facts = load_materials_facts()
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    arch = build_architecture_diagram()
    rag = build_rag_pipeline_diagram()

    doc = Document()
    configure_doc(doc, facts)

    add_title_page(doc, facts)
    add_contents(doc)

    add_section_title(doc, "ВВЕДЕНИЕ")
    add_paragraphs(doc, intro_paragraphs(facts))

    add_section_title(doc, "1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ АНАЛИЗА ЭКОНОМИЧЕСКИХ НОВОСТЕЙ")
    add_paragraphs(doc, theory_paragraphs())
    add_table(
        doc,
        ["Понятие", "Значение для работы"],
        [
            ["NLP", "Преобразует неструктурированный текст новостей в признаки для поиска и классификации."],
            ["Классификация влияния", "Определяет направленность новости: positive, neutral или negative."],
            ["Retrieval", "Находит релевантные источники по вопросу пользователя."],
            ["RAG", "Формирует ответ на основе найденного контекста, а не только параметров языковой модели."],
            ["Внешний LLM API", "Обобщает найденные документы и возвращает связный аналитический ответ."],
        ],
        widths=[4.0, 12.0],
    )

    add_section_title(doc, "2 ДАННЫЕ И ПОСТАНОВКА ML-ЗАДАЧИ")
    add_paragraphs(doc, data_paragraphs(facts))
    add_table(
        doc,
        ["Элемент постановки", "Описание"],
        [
            ["Объект", "Экономическая новость как текстовый документ."],
            ["Целевая переменная", "Класс влияния: positive, neutral, negative."],
            ["Вход модели", "Очищенный текст новости или embedding-представление."],
            ["Выход модели", "Предсказанный класс, confidence score и объяснение для интерфейса."],
            ["Метрики", "Accuracy и macro F1 на тестовом наборе."],
        ],
        widths=[4.5, 11.5],
    )

    add_section_title(doc, "3 МЕТОДЫ И МОДЕЛИ")
    add_paragraphs(doc, methods_paragraphs())
    add_table(
        doc,
        ["Модель", "Сильная сторона", "Ограничение"],
        [
            ["TF-IDF + Logistic Regression", "Высокая скорость, объяснимость, сильный baseline.", "Зависимость от лексических совпадений и словаря."],
            ["embedding-logreg", "Учитывает смысловую близость текстов.", "Может терять важные доменные маркеры текущей разметки."],
            ["tiny-transformer-classifier", "Проверяет нейросетевой подход.", "Ограничен размером модели и режимом обучения."],
        ],
        widths=[4.7, 5.7, 5.6],
    )

    add_section_title(doc, "4 АРХИТЕКТУРА ПРОГРАММНОЙ СИСТЕМЫ")
    add_paragraphs(doc, architecture_paragraphs())
    add_picture(doc, arch, "Схема 1 - Компактная архитектура программной системы")
    add_picture(doc, rag, "Схема 2 - ML/RAG-пайплайн обработки вопроса")

    add_section_title(doc, "5 РЕАЛИЗАЦИЯ И ПОЛЬЗОВАТЕЛЬСКИЙ СЦЕНАРИЙ")
    add_paragraphs(doc, realization_paragraphs())
    add_picture(doc, ASSETS_DIR / "chat-page.png", "Рисунок 1 - Страница чата", width=6.2)
    add_picture(doc, ASSETS_DIR / "ml-report-page.png", "Рисунок 2 - Страница ML-отчета", width=6.2)
    add_picture(doc, ASSETS_DIR / "topic-forecast-page.png", "Рисунок 3 - Страница прогноза", width=6.2)

    add_section_title(doc, "6 ЭКСПЕРИМЕНТ И РЕЗУЛЬТАТЫ")
    add_paragraphs(doc, experiment_paragraphs(facts))
    add_model_results_table(doc, facts)
    doc.add_paragraph(
        "Из таблицы видно, что TF-IDF + Logistic Regression является наиболее рациональным выбором для текущей версии системы. Она сочетает лучшее качество, простую интерпретацию и достаточную скорость для включения в пользовательский RAG-сценарий."
    )

    add_section_title(doc, "7 НАПРАВЛЕНИЯ УЛУЧШЕНИЯ")
    add_paragraphs(doc, improvements_paragraphs())
    add_bullets(
        doc,
        [
            "сбалансировать классы и проверить влияние весов классов на macro F1;",
            "уточнить правила разметки спорных экономических новостей;",
            "подобрать гиперпараметры для TF-IDF, embedding и transformer-подходов;",
            "проверить более сильную transformer-модель на том же test split;",
            "расширить признаки за счет тем, источников, дат и макроэкономических категорий;",
            "добавить временную динамику в тематическое прогнозирование.",
        ],
    )

    add_section_title(doc, "ЗАКЛЮЧЕНИЕ")
    add_paragraphs(doc, conclusion_paragraphs())

    add_sources(doc, facts)
    add_appendices(doc, facts)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    result = build_docx()
    print(result)
